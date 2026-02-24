import os
import sys
import csv
import argparse
import logging
import hashlib
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Iterable
from datetime import datetime

import chardet
import markdown
from bs4 import BeautifulSoup
from tqdm import tqdm

import PyPDF2
import docx


# ----------------------------
# Data Models
# ----------------------------

@dataclass
class Document:
    id: str
    text: str
    metadata: Dict[str, any]


@dataclass
class Chunk:
    id: str
    text: str
    metadata: Dict[str, any]


# ----------------------------
# Utility Functions
# ----------------------------

def detect_encoding(file_path: str) -> str:
    with open(file_path, "rb") as f:
        raw = f.read(10000)
    result = chardet.detect(raw)
    return result.get("encoding") or "utf-8"


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def safe_read_text(file_path: str) -> str:
    encoding = detect_encoding(file_path)
    with open(file_path, "r", encoding=encoding, errors="replace") as f:
        return f.read()


# ----------------------------
# Loader Implementations
# ----------------------------

class BaseLoader:
    def load(self, file_path: str) -> List[Document]:
        raise NotImplementedError


class PDFLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                documents.append(Document(
                    id=f"{file_path}-page-{i+1}",
                    text=text,
                    metadata={
                        "source": file_path,
                        "page": i + 1,
                        "file_type": "pdf"
                    }
                ))
        return documents


class DOCXLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return [Document(
            id=file_path,
            text=text,
            metadata={"source": file_path, "file_type": "docx"}
        )]


class TXTLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        text = safe_read_text(file_path)
        return [Document(
            id=file_path,
            text=text,
            metadata={"source": file_path, "file_type": "txt"}
        )]


class MarkdownLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        raw = safe_read_text(file_path)
        html = markdown.markdown(raw)
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text(separator="\n")
        return [Document(
            id=file_path,
            text=text,
            metadata={"source": file_path, "file_type": "markdown"}
        )]


class HTMLLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        raw = safe_read_text(file_path)
        soup = BeautifulSoup(raw, "html.parser")
        text = soup.get_text(separator="\n")
        return [Document(
            id=file_path,
            text=text,
            metadata={"source": file_path, "file_type": "html"}
        )]


class CSVLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        encoding = detect_encoding(file_path)
        rows = []
        with open(file_path, newline="", encoding=encoding, errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        text = "\n".join(rows)
        return [Document(
            id=file_path,
            text=text,
            metadata={"source": file_path, "file_type": "csv"}
        )]


class CodeLoader(BaseLoader):
    def load(self, file_path: str) -> List[Document]:
        text = safe_read_text(file_path)
        return [Document(
            id=file_path,
            text=text,
            metadata={"source": file_path, "file_type": "code"}
        )]


# ----------------------------
# Chunking Strategies
# ----------------------------

class ChunkingStrategy:
    def chunk(self, document: Document) -> List[Chunk]:
        raise NotImplementedError


class FixedSizeChunking(ChunkingStrategy):
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, document: Document) -> List[Chunk]:
        text = document.text
        chunks = []
        start = 0
        index = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            chunk_id = f"{document.id}-chunk-{index}"
            chunks.append(Chunk(
                id=chunk_id,
                text=chunk_text,
                metadata={**document.metadata, "chunk_index": index}
            ))
            start = end - self.overlap
            index += 1

        return chunks


class RecursiveChunking(ChunkingStrategy):
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, document: Document) -> List[Chunk]:
        paragraphs = document.text.split("\n\n")
        chunks = []
        current = ""
        index = 0

        for para in paragraphs:
            if len(current) + len(para) < self.chunk_size:
                current += para + "\n\n"
            else:
                chunks.append(self._make_chunk(document, current, index))
                index += 1
                current = para + "\n\n"

        if current.strip():
            chunks.append(self._make_chunk(document, current, index))

        return self._apply_overlap(chunks)

    def _make_chunk(self, document: Document, text: str, index: int) -> Chunk:
        return Chunk(
            id=f"{document.id}-chunk-{index}",
            text=text.strip(),
            metadata={**document.metadata, "chunk_index": index}
        )

    def _apply_overlap(self, chunks: List[Chunk]) -> List[Chunk]:
        if self.overlap <= 0:
            return chunks

        overlapped = []
        for i, chunk in enumerate(chunks):
            if i == 0:
                overlapped.append(chunk)
                continue

            prev = overlapped[-1]
            overlap_text = prev.text[-self.overlap:]
            new_text = overlap_text + "\n" + chunk.text
            overlapped.append(Chunk(
                id=chunk.id,
                text=new_text,
                metadata=chunk.metadata
            ))

        return overlapped


class SemanticChunking(ChunkingStrategy):
    def __init__(self, max_sentences: int = 5):
        self.max_sentences = max_sentences

    def chunk(self, document: Document) -> List[Chunk]:
        sentences = self._split_sentences(document.text)
        chunks = []
        buffer = []
        index = 0

        for sentence in sentences:
            buffer.append(sentence)
            if len(buffer) >= self.max_sentences:
                text = " ".join(buffer)
                chunks.append(Chunk(
                    id=f"{document.id}-chunk-{index}",
                    text=text,
                    metadata={**document.metadata, "chunk_index": index}
                ))
                buffer = []
                index += 1

        if buffer:
            text = " ".join(buffer)
            chunks.append(Chunk(
                id=f"{document.id}-chunk-{index}",
                text=text,
                metadata={**document.metadata, "chunk_index": index}
            ))

        return chunks

    def _split_sentences(self, text: str) -> List[str]:
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]


# ----------------------------
# Document Processor
# ----------------------------

class DocumentProcessor:
    def __init__(self, chunking_strategy: ChunkingStrategy):
        self.chunking_strategy = chunking_strategy
        self.loaders = {
            ".pdf": PDFLoader(),
            ".docx": DOCXLoader(),
            ".txt": TXTLoader(),
            ".md": MarkdownLoader(),
            ".html": HTMLLoader(),
            ".htm": HTMLLoader(),
            ".csv": CSVLoader(),
        }
        self.code_extensions = {
            ".py", ".js", ".ts", ".java", ".go", ".rs", ".cpp", ".c", ".h",
            ".json", ".yaml", ".yml", ".toml", ".sh"
        }

    def process_file(self, file_path: str) -> List[Chunk]:
        ext = os.path.splitext(file_path)[1].lower()

        loader = self.loaders.get(ext)
        if not loader and ext in self.code_extensions:
            loader = CodeLoader()

        if not loader:
            logging.warning(f"Unsupported file type: {file_path}")
            return []

        try:
            documents = loader.load(file_path)
        except Exception as e:
            logging.error(f"Failed to load {file_path}: {e}")
            return []

        chunks = []
        for doc in documents:
            try:
                doc_chunks = self.chunking_strategy.chunk(doc)
                chunks.extend(doc_chunks)
            except Exception as e:
                logging.error(f"Chunking failed for {doc.id}: {e}")

        return chunks

    def process_directory(self, input_dir: str) -> List[Chunk]:
        all_chunks = []
        files = self._collect_files(input_dir)

        for file_path in tqdm(files, desc="Processing documents"):
            chunks = self.process_file(file_path)
            all_chunks.extend(chunks)

        return all_chunks

    def _collect_files(self, input_dir: str) -> List[str]:
        collected = []
        for root, _, files in os.walk(input_dir):
            for name in files:
                collected.append(os.path.join(root, name))
        return collected


# ----------------------------
# CLI Interface
# ----------------------------

def build_chunking_strategy(strategy_name: str, chunk_size: int, overlap: int):
    if strategy_name == "fixed":
        return FixedSizeChunking(chunk_size=chunk_size, overlap=overlap)
    if strategy_name == "recursive":
        return RecursiveChunking(chunk_size=chunk_size, overlap=overlap)
    if strategy_name == "semantic":
        return SemanticChunking()
    raise ValueError(f"Unknown chunking strategy: {strategy_name}")


def main():
    parser = argparse.ArgumentParser(description="Universal Document Processor")
    parser.add_argument("--input", required=True, help="Input directory of documents")
    parser.add_argument("--strategy", default="recursive", choices=["fixed", "recursive", "semantic"])
    parser.add_argument("--chunk-size", type=int, default=500)
    parser.add_argument("--overlap", type=int, default=50)
    parser.add_argument("--output", default="processed_chunks.json")
    parser.add_argument("--log-level", default="INFO")

    args = parser.parse_args()

    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(asctime)s - %(levelname)s - %(message)s"
    )

    strategy = build_chunking_strategy(args.strategy, args.chunk_size, args.overlap)
    processor = DocumentProcessor(chunking_strategy=strategy)

    chunks = processor.process_directory(args.input)

    output_path = args.output
    export_chunks(chunks, output_path)

    logging.info(f"Processed {len(chunks)} chunks")
    logging.info(f"Saved to {output_path}")


# ----------------------------
# Export
# ----------------------------

def export_chunks(chunks: List[Chunk], output_path: str):
    import json

    serialized = []
    for chunk in chunks:
        serialized.append({
            "id": chunk.id,
            "text": chunk.text,
            "metadata": chunk.metadata
        })

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(serialized, f, ensure_ascii=False, indent=2)


# ----------------------------
# Entry Point
# ----------------------------

if __name__ == "__main__":
    main()
